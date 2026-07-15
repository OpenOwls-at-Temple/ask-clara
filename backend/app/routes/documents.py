import io
import logging

from bson import ObjectId
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.concurrency import run_in_threadpool
from fastapi.responses import StreamingResponse
from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.auth import get_current_user
from app.database import get_db, get_mongo_db
from app.documents.resumes import (
    get_generated_resumes_for_user,
    update_resume_edited_text,
)
from app.models.user import User
from app.schemas.resume import ResumeEditRequest, ResumeOut
from app.services import assessment_service
from app.services.resume_pdf import render_resume_pdf, render_resume_png

logger = logging.getLogger(__name__)

LLM_GENERATION_CAP = 20

router = APIRouter()


@router.post("/resumes/generate", response_model=list[ResumeOut])
async def generate_resumes(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    cap = 999999 if settings.environment == "local" else LLM_GENERATION_CAP
    quota_stmt = text(
        "UPDATE users SET llm_generation_count = llm_generation_count + 1 "
        "WHERE id = CAST(:user_id AS uuid) AND llm_generation_count < :cap "
        "RETURNING id"
    )
    quota_result = await db.execute(quota_stmt, {"user_id": str(user.id), "cap": cap})
    await db.commit()

    if quota_result.rowcount == 0:
        raise HTTPException(
            status_code=status.HTTP_429_TOO_MANY_REQUESTS,
            detail=(
                "You have reached the generation limit for this pilot. "
                "Please contact the team if you need more."
            ),
        )

    mongo = get_mongo_db()
    try:
        docs = await assessment_service.generate_resumes(db, mongo, user.id)
    except Exception as exc:
        await db.execute(
            text(
                "UPDATE users SET llm_generation_count = llm_generation_count - 1 "
                "WHERE id = CAST(:user_id AS uuid)"
            ),
            {"user_id": str(user.id)},
        )
        await db.commit()
        if isinstance(exc, ValueError):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc)
            )
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail=str(exc)
        )

    return [ResumeOut(**doc) for doc in docs]


@router.get("/resumes", response_model=list[ResumeOut])
async def list_resumes(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    mongo = get_mongo_db()
    docs = await get_generated_resumes_for_user(mongo, str(user.id))
    return [ResumeOut(**doc) for doc in docs]


@router.patch("/resumes/{resume_id}")
async def update_resume(
    resume_id: str,
    body: ResumeEditRequest,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    mongo = get_mongo_db()
    updated = await update_resume_edited_text(
        mongo, resume_id, str(user.id), body.edited_text
    )
    if not updated:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Resume not found."
        )
    return {"ok": True}


@router.get("/resumes/{resume_id}/download")
async def download_resume(
    resume_id: str,
    format: str = "pdf",
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Download a generated resume as a one-page Typst PDF (default) or DOCX.
    `format=png` returns an inline image of the same render, used by the UI
    as a chrome-free preview."""
    if format not in ("pdf", "docx", "png"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="format must be 'pdf', 'docx', or 'png'.",
        )

    mongo = get_mongo_db()
    try:
        doc = await mongo["resumes"].find_one({"_id": ObjectId(resume_id)})
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Resume not found."
        )

    if not doc or doc.get("user_id") != str(user.id):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Resume not found."
        )

    slug = (doc.get("target_title") or "resume").lower().replace(" ", "-")

    if format in ("pdf", "png"):
        renderer = render_resume_pdf if format == "pdf" else render_resume_png
        try:
            rendered = await run_in_threadpool(
                renderer,
                user.display_name or "Resume",
                doc.get("sections", []),
                doc.get("edited_text"),
            )
        except Exception:
            logger.exception("Typst resume rendering failed")
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail=(
                    "PDF generation failed — download the DOCX or copy the "
                    "text instead."
                ),
            )
        if format == "png":
            return StreamingResponse(
                io.BytesIO(rendered),
                media_type="image/png",
                headers={"Content-Disposition": "inline"},
            )
        return StreamingResponse(
            io.BytesIO(rendered),
            media_type="application/pdf",
            headers={
                "Content-Disposition": (
                    f'attachment; filename="clara-resume-{slug}.pdf"'
                )
            },
        )

    stream = await run_in_threadpool(_build_docx, doc, user.display_name or "Resume")
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="clara-resume-{slug}.docx"'
        },
    )


def _build_docx(doc: dict, student_name: str) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt

    document = Document()

    # Heading is the student's name — never the target job title.
    document.add_heading(student_name, level=0)

    edited_text = doc.get("edited_text")
    if edited_text:
        for line in edited_text.splitlines():
            document.add_paragraph(line)
    else:
        for section in doc.get("sections", []):
            heading = section.get("heading", "")
            content = section.get("content", "")
            if heading:
                document.add_heading(heading, level=2)
            if content:
                document.add_paragraph(content)

    stream = io.BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream

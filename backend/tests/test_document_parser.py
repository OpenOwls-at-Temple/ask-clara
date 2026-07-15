import io

import pytest

from app.services.document_parser_service import (
    UnsupportedFileTypeError,
    parse_document,
)

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def test_parse_document_extracts_docx_text():
    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Software Engineering Intern")
    doc.add_paragraph("Built dashboards in React")
    doc.save(buf)

    text = parse_document(buf.getvalue(), _DOCX_MIME, "resume.docx")

    assert "Software Engineering Intern" in text
    assert "Built dashboards in React" in text


def test_parse_document_reads_pdf():
    from pypdf import PdfWriter

    buf = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.write(buf)

    text = parse_document(buf.getvalue(), "application/pdf", "resume.pdf")

    assert text == ""  # blank page, but parsed without error


def test_parse_document_falls_back_to_extension():
    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("hello")
    doc.save(buf)

    assert "hello" in parse_document(buf.getvalue(), "", "resume.docx")


def test_parse_document_rejects_unsupported_type():
    with pytest.raises(UnsupportedFileTypeError):
        parse_document(b"plain text", "text/plain", "notes.txt")
